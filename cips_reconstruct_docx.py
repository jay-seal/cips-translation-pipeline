"""
cips_reconstruct_docx.py
========================
CIPS Translation Pipeline — DOCX Reconstruction Script

Triggered by GitHub Actions (reconstruct.yml or a dedicated docx_reconstruct.yml).
Accepts file paths via command-line arguments.

Usage:
    python cips_reconstruct_docx.py \
        --source  inputs/source.docx \
        --translations  inputs/agent3_output.json \
        [--qa  inputs/agent4_output.json] \
        --output  outputs/<filename>.docx \
        [--failure-threshold 0.10]

Outputs written to the same directory as --output:
    <name>.docx                     Translated document.
    <name>_match_report.json        Full per-segment match results.
    <name>_manual_corrections.html  Human-readable list of unmatched segments.

Matching strategy — applied in order per segment, stopping at first success:

    Tier 1   Exact normalised match on full paragraph text (body paragraphs).
    Tier 2   Substring normalised match (single-run paragraphs only, len > 3).
    Tier H   Header paragraphs — Tiers 1–2 applied to header text frames.
    Tier F   Footer paragraphs — Tiers 1–2 applied to footer text frames.

Normalisation (comparison only — never applied to output text):
    1. NFKC unicode normalisation.
    2. Collapse all whitespace to a single ASCII space.
    3. Strip leading/trailing whitespace.
    4. Strip leading bullet/list characters.

Run-level formatting preservation:
    Translated text is written to the first non-empty run of the paragraph.
    All other runs in the paragraph are cleared. This preserves the formatting
    (font, size, colour, bold, italic) from the first run across the full
    translated paragraph. For plain-body documents this is correct.
    For paragraphs with mixed intra-run formatting (e.g. one word bolded
    within a sentence), the translated text inherits the formatting of the
    first run. This is a known limitation and a backlog item.

Footer field elements:
    Word page number fields use <w:fldChar> / <w:instrText> run sequences.
    These are preserved intact — the reconstruction replaces only the
    surrounding text runs, leaving field elements untouched.

Exit codes:
    0   Success.
    1   Missing or invalid input files.
    2   Match failure rate exceeds --failure-threshold.
"""

import argparse
import html
import json
import logging
import re
import sys
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("cips_reconstruct_docx")

# ---------------------------------------------------------------------------
# Bullet/list character stripping
# ---------------------------------------------------------------------------
_LEADING_BULLET_RE = re.compile(
    '['
    '\u2022\u2023\u25E6\u2043\u2219\u25CF\u25AA'
    '\u2776-\u277F\u2780-\u2789'
    ']+' + r'\s*'
)

# ---------------------------------------------------------------------------
# Element types that are always processed regardless of is_in_text_box.
# (DOCX has no slide-layout/master concept, but we retain the guard for
# label elements that may be image-embedded captions.)
# ---------------------------------------------------------------------------
_ALWAYS_PROCESS_TYPES = frozenset({
    'heading', 'body_text', 'bullet_point', 'table_cell',
    'text_box', 'caption', 'paragraph',
})


# ---------------------------------------------------------------------------
# Normalisation
# ---------------------------------------------------------------------------

def normalise(text: str) -> str:
    """
    Normalise for comparison only. Never applied to output text.
    1. NFKC unicode normalisation.
    2. Smart quote / typographic punctuation normalisation — converts curly
       quotes and em/en dashes to their plain ASCII equivalents so that text
       extracted by AI agents (which may produce straight quotes) matches
       DOCX content (which typically uses Word's smart quote characters).
    3. Collapse whitespace to single space.
    4. Strip edges.
    5. Strip leading bullet/list characters.
    """
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    # Smart single quotes -> straight apostrophe
    text = text.replace('‘', "'").replace('’', "'")
    # Smart double quotes -> straight double quote
    text = text.replace('“', '"').replace('”', '"')
    # En dash / em dash -> hyphen
    text = text.replace('–', '-').replace('—', '-')
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    text = _LEADING_BULLET_RE.sub('', text)
    return text
def _has_field_elements(para) -> bool:
    """True if the paragraph contains Word field elements (e.g. PAGE fields)."""
    return bool(para._p.findall('.//' + qn('w:fldChar')))


# ---------------------------------------------------------------------------
# Run-level text replacement
# ---------------------------------------------------------------------------

def _replace_para_text(para, new_text: str) -> bool:
    """
    Write new_text into a paragraph preserving the first run's formatting.

    Returns True if replacement was made, False if no writable runs exist.

    Paragraphs containing field elements (page number fields etc.) are
    handled carefully: only text runs that are NOT part of a field sequence
    are modified. Field sequences are left intact.
    """
    runs = para.runs
    if not runs:
        return False

    if _has_field_elements(para):
        # For field-bearing paragraphs (e.g. footers with PAGE field),
        # find the plain text runs that precede the field sequence and
        # replace only those, leaving field runs untouched.
        plain_runs = [
            r for r in runs
            if r.text and not r._r.findall('.//' + qn('w:fldChar'))
               and not r._r.findall('.//' + qn('w:instrText'))
        ]
        if plain_runs:
            plain_runs[0].text = new_text
            for r in plain_runs[1:]:
                r.text = ""
            return True
        return False

    # Standard case: write to run[0], clear the rest.
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


# ---------------------------------------------------------------------------
# Single-paragraph matching
# ---------------------------------------------------------------------------

def _match_para(para, norm_source: str, translated_text: str,
                segment_id: str, tier_label: str) -> dict | None:
    """
    Apply Tier 1 (exact) and Tier 2 (substring) to a single paragraph.
    Returns a result dict on match, else None.
    """
    norm_para = normalise(para.text)
    if not norm_para:
        return None

    # Tier 1 — exact match
    if norm_para == norm_source:
        if _replace_para_text(para, translated_text):
            log.info("%-20s seg=%-12s  para=%s",
                     f"{tier_label}T1_EXACT", segment_id,
                     repr(norm_para[:50]))
            return {"result": f"{tier_label}T1_EXACT"}

    # Tier 2 — substring match (conservative: source must be >= 80% of
    # paragraph length to avoid silently overwriting longer paragraphs that
    # contain multiple independent sentences merged by Agent 1 into one).
    fraction = len(norm_source) / len(norm_para) if norm_para else 0
    if (len(norm_source) > 3
            and norm_source in norm_para
            and fraction >= 0.80):
        if _replace_para_text(para, translated_text):
            log.info("%-20s seg=%-12s  para=%s",
                     f"{tier_label}T2_SUBSTR", segment_id,
                     repr(norm_para[:50]))
            return {"result": f"{tier_label}T2_SUBSTRING"}

    return None


# ---------------------------------------------------------------------------
# Core reconstruction
# ---------------------------------------------------------------------------

def reconstruct(docx_path: Path, json_path: Path, output_path: Path) -> dict:
    """Load, translate, and save the document. Returns a match report dict."""
    log.info("Loading document: %s", docx_path)
    doc = Document(str(docx_path))

    log.info("Loading translation JSON: %s", json_path)
    with json_path.open(encoding="utf-8") as fh:
        data = json.load(fh)

    segments = data.get("segments", [])
    log.info("Total segments in JSON: %d", len(segments))

    # Collect all addressable paragraphs:
    # body paragraphs + header paragraphs + footer paragraphs
    body_paras = list(doc.paragraphs)

    header_paras = []
    footer_paras = []
    for section in doc.sections:
        header_paras.extend(section.header.paragraphs)
        footer_paras.extend(section.footer.paragraphs)

    # Also collect table cell paragraphs
    table_paras = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_paras.extend(cell.paragraphs)

    log.info("Addressable paragraphs — body: %d, headers: %d, footers: %d, table cells: %d",
             len(body_paras), len(header_paras), len(footer_paras), len(table_paras))

    results = []
    counts = {
        "total": 0, "translated": 0,
        "skipped_no_translation": 0, "skipped_do_not_translate": 0,
        "skipped_not_text_box": 0,
        "T1_EXACT": 0, "T2_SUBSTRING": 0,
        "HEADER_T1_EXACT": 0, "HEADER_T2_SUBSTRING": 0,
        "FOOTER_T1_EXACT": 0, "FOOTER_T2_SUBSTRING": 0,
        "TABLE_T1_EXACT": 0, "TABLE_T2_SUBSTRING": 0,
        "NO_MATCH": 0, "SKIP_EMPTY_SOURCE": 0,
    }

    for seg in segments:
        counts["total"] += 1
        segment_id   = seg.get("segment_id", "UNKNOWN")
        source_text  = seg.get("source_text", "")
        translated   = seg.get("translated_text")
        status       = seg.get("translation_status", "")
        element_type = seg.get("element_type", "")
        page_number  = seg.get("slide_or_page")

        if not translated:
            counts["skipped_no_translation"] += 1
            continue
        if status == "DO_NOT_TRANSLATE":
            counts["skipped_do_not_translate"] += 1
            continue

        # is_in_text_box guard — same logic as PPTX script
        if seg.get("is_in_text_box") is False:
            if element_type not in _ALWAYS_PROCESS_TYPES:
                log.debug("SKIP_NOT_TEXT_BOX  seg=%-12s  source=%r",
                          segment_id, source_text[:60])
                counts["skipped_not_text_box"] += 1
                results.append({
                    "segment_id": segment_id,
                    "result": "SKIP_NOT_TEXT_BOX",
                    "page": page_number,
                })
                continue

        norm_source = normalise(source_text)
        if not norm_source:
            counts["SKIP_EMPTY_SOURCE"] += 1
            results.append({"segment_id": segment_id,
                            "result": "SKIP_EMPTY_SOURCE", "page": page_number})
            continue

        counts["translated"] += 1
        outcome = None

        # --- Body paragraphs ---
        for para in body_paras:
            outcome = _match_para(para, norm_source, translated,
                                  segment_id, "")
            if outcome:
                break

        # --- Table cells ---
        if not outcome:
            for para in table_paras:
                outcome = _match_para(para, norm_source, translated,
                                      segment_id, "TABLE_")
                if outcome:
                    break

        # --- Header ---
        if not outcome:
            for para in header_paras:
                outcome = _match_para(para, norm_source, translated,
                                      segment_id, "HEADER_")
                if outcome:
                    break

        # --- Footer ---
        if not outcome:
            for para in footer_paras:
                outcome = _match_para(para, norm_source, translated,
                                      segment_id, "FOOTER_")
                if outcome:
                    break

        if outcome:
            outcome["segment_id"] = segment_id
            outcome["page"] = page_number
            result_key = outcome["result"]
            counts[result_key] = counts.get(result_key, 0) + 1
            results.append(outcome)
        else:
            log.warning("NO MATCH  seg=%-12s  source=%r",
                        segment_id, source_text[:80])
            counts["NO_MATCH"] += 1
            results.append({
                "segment_id": segment_id,
                "result": "NO_MATCH",
                "source_text": source_text,
                "translated_text": translated,
                "page": page_number,
            })

    output_path.parent.mkdir(parents=True, exist_ok=True)
    log.info("Saving translated document to %s", output_path)
    doc.save(str(output_path))
    log.info("Save complete.")

    t1_total = counts["T1_EXACT"] + counts["HEADER_T1_EXACT"] + counts["FOOTER_T1_EXACT"] + counts["TABLE_T1_EXACT"]
    t2_total = counts["T2_SUBSTRING"] + counts["HEADER_T2_SUBSTRING"] + counts["FOOTER_T2_SUBSTRING"] + counts["TABLE_T2_SUBSTRING"]
    total_handled = t1_total + t2_total
    no_match = counts["NO_MATCH"]
    handle_pct = (total_handled / counts["translated"] * 100) if counts["translated"] else 0
    fail_rate = no_match / counts["translated"] if counts["translated"] else 0

    log.info("=" * 60)
    log.info("RECONSTRUCTION SUMMARY")
    log.info("  Total segments              : %d", counts["total"])
    log.info("  Segments with translation   : %d", counts["translated"])
    log.info("  Body T1 exact               : %d", counts["T1_EXACT"])
    log.info("  Body T2 substring           : %d", counts["T2_SUBSTRING"])
    log.info("  Table T1 exact              : %d", counts["TABLE_T1_EXACT"])
    log.info("  Header T1 exact             : %d", counts["HEADER_T1_EXACT"])
    log.info("  Footer T1 exact             : %d", counts["FOOTER_T1_EXACT"])
    log.info("  Footer T2 substring         : %d", counts["FOOTER_T2_SUBSTRING"])
    log.info("  Total handled               : %d  (%.1f%%)", total_handled, handle_pct)
    log.info("  No match (manual required)  : %d  (%.1f%%)", no_match, fail_rate * 100)
    log.info("  Skipped (no translation)    : %d", counts["skipped_no_translation"])
    log.info("  Skipped (not text box)      : %d", counts["skipped_not_text_box"])
    log.info("=" * 60)

    return {
        "summary": counts,
        "handle_rate_pct": round(handle_pct, 2),
        "no_match_rate_pct": round(fail_rate * 100, 2),
        "results": results,
        "no_match_segments": [r for r in results if r.get("result") == "NO_MATCH"],
    }


# ---------------------------------------------------------------------------
# HTML corrections report
# ---------------------------------------------------------------------------

def _write_html_report(report: dict, output_path: Path,
                       source_filename: str) -> None:
    no_match = report.get("no_match_segments", [])
    summary  = report.get("summary", {})
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    rows_html = ""
    for item in no_match:
        page = item.get("page", "—")
        seg_id = item.get("segment_id", "—")
        src = item.get("source_text", "")
        translated = item.get("translated_text", "")
        rows_html += f"""
        <tr>
          <td class="page">{html.escape(str(page))}</td>
          <td class="seg">{html.escape(seg_id)}</td>
          <td class="src">{html.escape(src).replace(chr(10), '<br>')}</td>
          <td class="tgt">{html.escape(translated).replace(chr(10), '<br>')}</td>
        </tr>"""

    total = summary.get("translated", 0)
    no_match_n = summary.get("NO_MATCH", 0)
    handled_n = total - no_match_n

    page = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Manual Corrections — {html.escape(source_filename)}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            font-size: 14px; color: #1a1a2e; background: #f4f6fb; padding: 32px 24px; }}
    header {{ max-width: 1100px; margin: 0 auto 28px; }}
    h1 {{ font-size: 22px; font-weight: 700; margin-bottom: 6px; }}
    .meta {{ color: #555; font-size: 13px; margin-bottom: 18px; }}
    .stats {{ display: flex; gap: 16px; flex-wrap: wrap; margin-bottom: 28px; }}
    .stat {{ background: #fff; border: 1px solid #dde3f0; border-radius: 8px;
             padding: 14px 20px; min-width: 140px; }}
    .stat .label {{ font-size: 11px; text-transform: uppercase; letter-spacing: .06em;
                   color: #777; margin-bottom: 4px; }}
    .stat .value {{ font-size: 26px; font-weight: 700; }}
    .stat.warn .value {{ color: #c0392b; }}
    .stat.ok   .value {{ color: #27ae60; }}
    .table-wrap {{ max-width: 1100px; margin: 0 auto; overflow-x: auto; }}
    table {{ width: 100%; border-collapse: collapse; background: #fff;
             border-radius: 8px; overflow: hidden;
             box-shadow: 0 1px 4px rgba(0,0,0,.08); }}
    thead th {{ background: #1a1a2e; color: #fff; font-size: 12px;
               text-transform: uppercase; letter-spacing: .06em;
               padding: 12px 14px; text-align: left; }}
    tbody tr:nth-child(odd)  {{ background: #fff; }}
    tbody tr:nth-child(even) {{ background: #f9fafc; }}
    tbody tr:hover {{ background: #eef2ff; }}
    td {{ padding: 10px 14px; vertical-align: top; line-height: 1.5;
         border-bottom: 1px solid #eee; }}
    td.page {{ font-weight: 700; width: 60px; text-align: center; }}
    td.seg  {{ font-family: monospace; font-size: 12px; color: #555; white-space: nowrap; }}
    td.src  {{ color: #333; }}
    td.tgt  {{ color: #1a6b3a; font-weight: 500; }}
    .empty  {{ text-align: center; padding: 40px; color: #888; font-size: 16px; }}
  </style>
</head>
<body>
  <header>
    <h1>Manual Corrections Required</h1>
    <p class="meta">Source: {html.escape(source_filename)} &nbsp;·&nbsp; Generated: {ts}</p>
    <div class="stats">
      <div class="stat ok"><div class="label">Auto-applied</div><div class="value">{handled_n}</div></div>
      <div class="stat warn"><div class="label">Needs manual fix</div><div class="value">{no_match_n}</div></div>
      <div class="stat"><div class="label">Total segments</div><div class="value">{total}</div></div>
    </div>
  </header>
  <div class="table-wrap">
    <table>
      <thead>
        <tr><th>Page</th><th>Segment</th>
            <th>English (source)</th><th>Translation (target)</th></tr>
      </thead>
      <tbody>
        {"<tr><td colspan='4' class='empty'>No manual corrections required.</td></tr>" if not no_match else rows_html}
      </tbody>
    </table>
  </div>
</body>
</html>"""

    output_path.write_text(page, encoding="utf-8")
    log.info("Manual corrections report written: %s  (%d items)",
             output_path, no_match_n)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="CIPS Translation Pipeline — DOCX Reconstruction Script"
    )
    parser.add_argument("--source", required=True, help="Path to the source DOCX.")
    parser.add_argument("--translations", required=True,
                        help="Path to the Agent 3 translation JSON.")
    parser.add_argument("--qa", required=False, default=None,
                        help="Path to the Agent 4 QA JSON (optional).")
    parser.add_argument("--output", required=True,
                        help="Destination path for the translated DOCX.")
    parser.add_argument("--failure-threshold", type=float, default=0.20,
                        help=(
                            "Exit with code 2 if the NO_MATCH rate exceeds "
                            "this fraction. Default 0.20 (20%%). Agent 1 may "
                            "split a single DOCX paragraph into multiple "
                            "segments; these produce NO_MATCH at reconstruction "
                            "time but are a structural ingestion issue rather "
                            "than inaccessible content. Set lower once paragraph "
                            "combining is implemented in the script."
                        ))
    args = parser.parse_args()

    source_path = Path(args.source)
    json_path   = Path(args.translations)
    output_path = Path(args.output)

    if not source_path.is_file():
        log.error("Source DOCX not found: %s", source_path)
        sys.exit(1)
    if not json_path.is_file():
        log.error("Translation JSON not found: %s", json_path)
        sys.exit(1)

    report = reconstruct(source_path, json_path, output_path)

    report_path = output_path.parent / output_path.name.replace(".docx", "_match_report.json")
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2),
                           encoding="utf-8")
    log.info("Match report written to %s", report_path)

    html_path = output_path.parent / output_path.name.replace(".docx", "_manual_corrections.html")
    _write_html_report(report, html_path, source_path.name)

    no_match_rate = (report["summary"].get("NO_MATCH", 0)
                     / max(report["summary"].get("translated", 1), 1))
    if no_match_rate > args.failure_threshold:
        log.error(
            "NO_MATCH rate %.1f%% exceeds threshold %.0f%%.",
            no_match_rate * 100, args.failure_threshold * 100,
        )
        sys.exit(2)

    log.info("Reconstruction complete. Output: %s", output_path)


if __name__ == "__main__":
    main()
